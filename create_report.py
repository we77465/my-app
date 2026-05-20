# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IMG_DIR = r'D:\work\上傳到公司\screenshots_ops'
OUT_FILE = r'D:\work\上傳到公司\CHBApp_功能操作測試報告.docx'

doc = Document()

# Page size: A4
section = doc.sections[0]
section.page_width  = int(8.27 * 914400)
section.page_height = int(11.69 * 914400)
section.left_margin   = int(1.0 * 914400)
section.right_margin  = int(1.0 * 914400)
section.top_margin    = int(1.0 * 914400)
section.bottom_margin = int(1.0 * 914400)

def add_image(doc, filename, width_inches=5.5):
    p = os.path.join(IMG_DIR, filename)
    if not os.path.exists(p):
        doc.add_paragraph(f'[圖片缺失: {filename}]')
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(p, width=Inches(width_inches))

def add_caption(doc, text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def h1(doc, text):
    doc.add_heading(text, level=1)

def h2(doc, text):
    doc.add_heading(text, level=2)

def body(doc, text):
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.size = Pt(12)

# --- Cover Page ---
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover.add_run('CHBApp.BK')
run.bold = True
run.font.size = Pt(28)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('員工薪資系統 功能操作測試報告')
r2.bold = True
r2.font.size = Pt(20)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.add_run('測試日期：2026/05/17').font.size = Pt(14)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.add_run('測試帳號：1111 / 1111').font.size = Pt(14)

doc.add_page_break()

# --- 1. Login ---
h1(doc, '1. 登入功能')
body(doc, '測試登入畫面，輸入帳號 1111 / 密碼 1111，驗證成功後進入主選單。')
add_image(doc, 'LOGIN_01_登入畫面.jpg')
add_caption(doc, '圖 1-1：登入畫面')
add_image(doc, 'LOGIN_02_主選單.jpg')
add_caption(doc, '圖 1-2：登入成功，主選單')
doc.add_page_break()

# --- 2. BK101 ---
h1(doc, '2. BK101 員工基本資料建檔')
body(doc, '測試新增員工資料功能。輸入員工編號、姓名、身份字號、帳號等欄位，儲存後查詢確認，並測試刪除功能。')
h2(doc, '2.1 空白畫面')
add_image(doc, 'BK101_01_空白畫面.jpg')
add_caption(doc, '圖 2-1：BK101 空白建檔畫面')
h2(doc, '2.2 填入資料')
add_image(doc, 'BK101_02_填資料.jpg')
add_caption(doc, '圖 2-2：填入測試員工資料')
h2(doc, '2.3 儲存成功')
add_image(doc, 'BK101_03_儲存成功.jpg')
add_caption(doc, '圖 2-3：儲存成功訊息')
h2(doc, '2.4 查詢結果')
add_image(doc, 'BK101_04_查詢結果.jpg')
add_caption(doc, '圖 2-4：查詢已建立之員工')
h2(doc, '2.5 刪除確認')
add_image(doc, 'BK101_05_刪除確認.jpg')
add_caption(doc, '圖 2-5：刪除確認對話框')
h2(doc, '2.6 刪除成功')
add_image(doc, 'BK101_06_刪除成功.jpg')
add_caption(doc, '圖 2-6：員工資料刪除完成')
doc.add_page_break()

# --- 3. BK102 ---
h1(doc, '3. BK102 員工資料查詢')
body(doc, '測試員工資料維護查詢功能，包含空白查詢（全部列表）及關鍵字查詢。')
h2(doc, '3.1 查詢畫面（無資料）')
add_image(doc, 'BK102_01_查詢畫面.jpg')
add_caption(doc, '圖 3-1：BK102 查詢畫面，初始 0 筆')
h2(doc, '3.2 查詢結果（有資料）')
add_image(doc, 'BK102_02_查詢結果.jpg')
add_caption(doc, '圖 3-2：BK102 查詢結果，共 1 筆 TEST01')
doc.add_page_break()

# --- 4. BK201 ---
h1(doc, '4. BK201 薪資輸入－個人')
body(doc, '測試個別員工薪資輸入功能。找到 TEST01，輸入本月薪資 50,000，儲存確認。')
h2(doc, '4.1 空白薪資畫面')
add_image(doc, 'BK201_01_空白.jpg')
add_caption(doc, '圖 4-1：BK201 個人薪資輸入，初始畫面')
h2(doc, '4.2 儲存成功')
add_image(doc, 'BK201_02_儲存成功.jpg')
add_caption(doc, '圖 4-2：薪資儲存成功，本月薪資 50,000')
doc.add_page_break()

# --- 5. BK202 ---
h1(doc, '5. BK202 薪資輸入－全公司')
body(doc, '測試全公司薪資統一輸入功能，可批次設定所有員工的本月薪資及存提區分。')
add_image(doc, 'BK202_01_全公司.jpg')
add_caption(doc, '圖 5-1：BK202 全公司薪資輸入畫面')
doc.add_page_break()

# --- 6. BK301 ---
h1(doc, '6. BK301 列印作業')
body(doc, '測試本月薪資清冊列印功能，包含列印設定及預覽。')
h2(doc, '6.1 列印設定')
add_image(doc, 'BK301_01_列印設定.jpg')
add_caption(doc, '圖 6-1：BK301 列印設定畫面')
h2(doc, '6.2 列印預覽')
add_image(doc, 'BK301_02_預覽.jpg')
add_caption(doc, '圖 6-2：本月薪資清冊列印預覽，TEST01 共 1 筆，合計 50,000')
doc.add_page_break()

# --- 7. BK401 ---
h1(doc, '7. BK401 轉出磁片作業－全體（百年規格）')
body(doc, '測試全體員工薪資轉出磁片功能（百年規格），設定撥帳日期、檔案路徑等參數。')
add_image(doc, 'BK401_01_全體磁片.jpg')
add_caption(doc, '圖 7-1：BK401 轉出磁片全體畫面')
doc.add_page_break()

# --- 8. BK4011 ---
h1(doc, '8. BK4011 轉出磁片作業－個別（百年規格）')
body(doc, '測試個別員工薪資轉出磁片功能（百年規格），可指定特定員工帳號輸出。')
add_image(doc, 'BK4011_01_個別磁片.jpg')
add_caption(doc, '圖 8-1：BK4011 轉出磁片個別畫面')
doc.add_page_break()

# --- 9. BK501 ---
h1(doc, '9. BK501 二次轉帳輸入（結果磁片匯入）')
body(doc, '測試二次轉扣帳作業，匯入銀行退件結果磁片，重新處理退件帳號。')
add_image(doc, 'BK501_01_二次轉帳.jpg')
add_caption(doc, '圖 9-1：BK501 薪資磁片匯入（二扣作業）畫面')
doc.add_page_break()

# --- 10. 密碼變更 ---
h1(doc, '10. 密碼變更')
body(doc, '測試密碼作業功能。點選密碼作業選單觸發密碼變更提示對話框，確認功能正常運作。')
add_image(doc, '密碼變更_01_提示.jpg')
add_caption(doc, '圖 10-1：密碼變更提示畫面')
doc.add_page_break()

# --- Summary Table ---
h1(doc, '測試總結')
body(doc, '本次對 CHBApp.BK 新版員工薪資管理系統（.NET 8）進行完整功能操作測試，共涵蓋 10 項功能模組，測試結果如下：')
doc.add_paragraph()

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '功能代號'
hdr[1].text = '功能名稱'
hdr[2].text = '測試結果'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows_data = [
    ('Login',    '登入',                    '✓ 通過'),
    ('BK101',   '員工基本資料建檔',          '✓ 通過'),
    ('BK102',   '員工資料查詢',              '✓ 通過'),
    ('BK201',   '薪資輸入－個人',            '✓ 通過'),
    ('BK202',   '薪資輸入－全公司',          '✓ 通過'),
    ('BK301',   '列印作業',                  '✓ 通過'),
    ('BK401',   '轉出磁片－全體（百年規格）','✓ 通過'),
    ('BK4011',  '轉出磁片－個別（百年規格）','✓ 通過'),
    ('BK501',   '二次轉帳輸入',              '✓ 通過'),
    ('密碼變更','密碼作業',                  '✓ 通過'),
]
for code, name, result in rows_data:
    row = table.add_row().cells
    row[0].text = code
    row[1].text = name
    row[2].text = result

doc.add_paragraph()
body(doc, '結論：CHBApp.BK 新版系統所有主要功能均運作正常，操作流程符合預期。')

doc.save(OUT_FILE)
print(f'Report saved to: {OUT_FILE}')
